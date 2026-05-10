import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import spacy

# Page configuration
st.set_page_config(page_title="CV Matcher", layout="wide")

# Title
st.title("AI CV Matcher")
st.markdown("Upload your CV and a job description")
  
# Sidebar for file uploads
st.sidebar.header("Upload Files")
cv_file = st.sidebar.file_uploader(
    "Upload your CV",
    type=["pdf", "docx", "txt"]
)

jd_file = st.sidebar.file_uploader(
    "Upload Job Description",
    type=["pdf", "docx", "txt"]
)
# File reading functions
def read_file(file):

    if file.name.endswith(".pdf"):
        reader = PdfReader(file)
        return "".join(
            [page.extract_text() or "" for page in reader.pages]
        )

    elif file.name.endswith(".docx"):
        doc = Document(file)
        return "\n".join(
            [p.text for p in doc.paragraphs]
        )

    elif file.name.endswith(".txt"):
        return file.read().decode("utf-8")

    else:
        return ""


    """def read_pdf(file):
        reader = PdfReader(file)
        return "".join([page.extract_text() or "" for page in reader.pages])

    def read_docx(file):
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])"""


# Load models
@st.cache_resource
def load_models():
    model = SentenceTransformer('all-MiniLM-L6-v2')
    nlp = spacy.load("en_core_web_sm")
    return model, nlp

model, nlp = load_models()

# Similarity function
def get_similarity(text1, text2):
    emb1 = model.encode(text1)
    emb2 = model.encode(text2)
    return cosine_similarity([emb1], [emb2])[0][0]

# Extract skills
def extract_skills(text):
    doc = nlp(text)
    skills = set()
    for token in doc:
        if token.pos_ in ["NOUN", "PROPN"]:
            skills.add(token.text.lower())
    return skills

# Additional skills input
additional_skills = st.text_area(
    "Enter additional skills (comma-separated)",
    placeholder="Python, AWS, Docker"
)

extra_skills = set()

if additional_skills:
    extra_skills = set(
        skill.strip().lower()
        for skill in additional_skills.split(",")
        if skill.strip()
    )


# Main analysis
if cv_file and jd_file:
    cv_text = read_file(cv_file)
    jd_text = read_file(jd_file)
    
    # Semantic similarity
    semantic_score = get_similarity(cv_text, jd_text)

    # Extract skills
    jd_skills = extract_skills(jd_text)
    cv_skills = extract_skills(cv_text)

# Add Manually Entered Skills
    jd_skills = jd_skills.union(extra_skills)

    #Compare Skills
    matched = jd_skills & cv_skills
    missing = jd_skills - cv_skills

    #Skill score calculation
    
    skill_score = len(matched) / len(jd_skills) if len(jd_skills) > 0 else 0
    overall = (0.6 * skill_score) + (0.4 * semantic_score)
    
    # Display results
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Overall Match", f"{overall*100:.1f}%")
    with col2:
        st.metric("Skill Match", f"{skill_score*100:.1f}%")
    with col3:
        st.metric("Semantic Match", f"{semantic_score*100:.1f}%")
    
    st.divider()
    
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("✅ Matched Skills")
        st.write(", ".join(sorted(matched)) if matched else "No skills matched")
    with col2:
        st.subheader("❌ Missing Skills")
        st.write(", ".join(sorted(missing)) if missing else "All skills present!")
else:
    st.info("👈 Please upload both files in the sidebar to begin analysis")
#END OF CODE
#Thank you
#TESTING

